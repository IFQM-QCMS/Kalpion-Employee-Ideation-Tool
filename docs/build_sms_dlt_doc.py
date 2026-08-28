# -*- coding: utf-8 -*-
"""
Build docs/SMS_DLT_TEMPLATES.docx — the SMS/DLT template request pack.

    python docs/build_sms_dlt_doc.py

Everything in it was read out of the running code, not invented:

    backend/src/config/index.js          config.sms.templates / config.sms.text
    backend/src/services/smsService.js   messageFor(), fillTemplate(), sendTestSms()
    backend/src/services/otpService.js   the sign-in code path
    backend/src/services/verificationService.js  PURPOSES
    backend/src/services/registrationService.js  registration_phone
    backend/src/services/authService.js          password_reset
    backend/src/services/userService.js          phone_verify + the change alert
    backend/.env                                 the ids and wording in force today

The document is meant to be sent to the operator as-is, so it leads with the
table they need and keeps the reasoning behind it.
"""
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, 'SMS_DLT_TEMPLATES.docx')

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x55, 0x55, 0x55)
RULE = RGBColor(0x99, 0x99, 0x99)

# ── Facts, as read from the code ────────────────────────────────────────────

PE_ID = '1201174858303838784'
HEADER = 'IFQMID-T'
ENTITY = 'IFQM Ideation'
SID = 'HXAP1678914824IN'

# Registered by Jio on 26 Aug 2026 under header IFQMID-T. These mirror
# backend/src/config/smsTemplates.js, which is the source the code reads; if the
# two ever disagree, the code is right and this document is stale.
#
#   (purpose, label, template id, approved wording)
REGISTERED = [
    ('registration_phone', 'Registration OTP', '1277178671564743852',
     'Dear Customer, use OTP {#number#} to complete your registration on '
     'IFQM Ideation. Do not share this OTP with anyone.'),
    ('login', 'Sign-in OTP', '1277178730169418603',
     'Dear Customer, use OTP {#number#} to complete your sign-in on '
     'IFQM Ideation. Do not share this OTP with anyone.'),
    ('password_reset', 'Password Reset OTP', '1277178730612100625',
     'Dear Customer, use OTP {#number#} to reset your password on '
     'IFQM Ideation. Do not share this OTP with anyone.'),
]

# Submitted, not yet granted. Jio classified it as Service Implicit rather than
# Transactional; a revised submission is in progress.
PENDING = (
    'phone_changed', 'Mobile Number Changed — Security Alert',
    'Your IFQM Ideation sign-in number was changed to one ending {#number#}. '
    'If this was not you, contact your administrator.',
    'Classified Service Implicit rather than Transactional. Awaiting '
    're-submission for Transactional approval.',
)

# Every place the platform sends an SMS.
USAGE = [
    ('1', 'Sign-in by one-time code',
     'login',
     'A person enters their email address or mobile number on the sign-in '
     'screen while one-time-code sign-in is switched on.',
     'otpService.requestOtp()',
     'Yes'),
    ('2', 'Company registration — verify the mobile number',
     'registration_phone',
     'A company fills in the self-registration form and asks for a code to '
     'confirm the mobile number on the application.',
     'registrationService (sendPhoneCode)',
     'Yes'),
    ('3', 'Password reset',
     'password_reset',
     'A person uses "Forgot password" and asks for the code by SMS rather '
     'than email.',
     'authService (requestPasswordReset)',
     'Yes'),
    ('4', 'Confirm a change of mobile number',
     'phone_verify',
     'A signed-in person edits the mobile number on My Profile; the code goes '
     'to the NEW number to prove they hold it.',
     'userService.requestPhoneChangeCode()',
     'Yes'),
    ('5', 'Alert the old number after a change',
     '(none — see 2.2)',
     'Immediately after a mobile number is changed, the PREVIOUS number is '
     'told, so the real owner learns about it if the change was not theirs.',
     'userService.notifyPhoneChanged()',
     'No'),
    ('6', 'Gateway test message',
     'login',
     'A platform administrator presses "Test Connection" on the Messaging '
     'screen. Sends the registered wording with a dummy code of 000000.',
     'smsService.sendTestSms()',
     'Yes'),
]

# What we are asking Jio DLT to register.
REQUESTS = [
    ('T1', 'Sign-in OTP', 'Service Implicit', '2',
     '{#var#} is your IFQM sign-in code. It expires in {#var#} minute(s). '
     'Do not share it with anyone.',
     'Journey 1 and the gateway test (6).'),
    ('T2', 'Registration / activation OTP', 'Service Implicit', '2',
     '{#var#} is your IFQM verification code. It expires in {#var#} minute(s). '
     'Do not share it with anyone.',
     'Journeys 2 and 4 — confirming a mobile number, at registration and on '
     'an existing account.'),
    ('T3', 'Password reset OTP', 'Service Implicit', '2',
     '{#var#} is your IFQM password reset code. It expires in {#var#} '
     'minute(s). Do not share it with anyone.',
     'Journey 3.'),
    ('T4', 'Mobile number changed — security alert', 'Service Implicit', '1',
     'Your IFQM sign-in number was changed to one ending {#var#}. If this was '
     'not you, contact your administrator.',
     'Journey 5. Not a one-time code — a security notice to the number that '
     'was replaced.'),
]


# ── Formatting helpers ──────────────────────────────────────────────────────

def style_doc(d):
    st = d.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.15
    for sec in d.sections:
        sec.left_margin = sec.right_margin = Inches(0.85)
        sec.top_margin = sec.bottom_margin = Inches(0.8)


def h(d, text, level=1, space_before=16):
    p = d.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = INK
        r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def para(d, text, bold=False, muted=False, size=10.5, justify=True, space_after=7):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = MUTED if muted else INK
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p


def mono(d, text, indent=0.18):
    """Message wording, set apart so the exact characters are unambiguous."""
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    _shade(p, 'F4F4F4')
    return p


def bullet(d, text, bold_prefix=None):
    p = d.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _shade(p, hex_fill):
    pr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_fill)
    pr.append(shd)


def _cell_shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def table(d, headers, rows, widths, font=9, mono_cols=()):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(font)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        _cell_shade(c, 'E8E8E8')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(font)
            if i in mono_cols:
                r.font.name = 'Consolas'
                r.font.size = Pt(8.5)
            p.paragraph_format.space_after = Pt(2)
            for j, w in enumerate(widths):
                cells[j].width = Inches(w)
    return t


# ── Document ────────────────────────────────────────────────────────────────

def build():
    d = Document()
    style_doc(d)

    # ── Cover ──
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Kalpion')
    r.bold = True
    r.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SMS Usage and DLT Content Template Request')
    r.bold = True
    r.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(4)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Prepared for registration with Jio TrueConnect (DLT)')
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(18)

    table(d, ['Field', 'Value'], [
        ['Principal Entity (PE) name', ENTITY],
        ['Principal Entity ID', PE_ID],
        ['Header / Sender ID', HEADER],
        ['Telemarketer / gateway', 'Kaleyra'],
        ['Kaleyra account SID', SID],
        ['Templates already registered', '3 (see section 4)'],
        ['Templates requested here', '4 (see section 3)'],
    ], [2.3, 4.3], font=10)

    d.add_paragraph()
    para(d, 'This document has two purposes. Sections 1 and 2 record exactly '
            'where the platform sends an SMS and what it sends, so the request '
            'can be checked against the product rather than taken on trust. '
            'Section 3 is the request itself, in the form the operator needs.',
         muted=True)

    d.add_page_break()

    # ── 1. Summary ──
    h(d, '1.  What is being asked for, in one page', 1, space_before=0)

    para(d, 'The platform sends an SMS in six situations. Five of them carry a '
            'one-time code; the sixth is a security notice sent to a mobile '
            'number that has just been replaced. Four content templates cover '
            'all six.')

    para(d, 'Three of those four are registered and in use. The fourth is '
            'submitted and not yet granted.')

    table(d,
          ['Journey', 'Template ID', 'Status'],
          [[label, tid, 'Registered'] for _p, label, tid, _t in REGISTERED]
          + [[PENDING[1], '—', 'Pending']],
          [2.3, 1.9, 1.6], font=9, mono_cols=(1,))

    d.add_paragraph()

    para(d, 'What the pending one costs, until it is granted.', bold=True)
    bullet(d, 'The security notice sent to a mobile number that has just been '
              'replaced (journey 5) is the only warning the rightful owner of a '
              'number gets if somebody else moves an account onto their own '
              'handset. It is not sent while it has no registration.',
           bold_prefix='Effect: ')
    bullet(d, 'It is not sent rather than sent-and-dropped. A message with no '
              'registration is accepted by the gateway and discarded by the '
              'carrier, with no error at either end — so sending it anyway '
              'would have the platform record a delivery that never happened. '
              'For a security alert that is worse than silence, because it '
              'reads as success. The e-mail alert still goes out.',
           bold_prefix='Handling: ')
    bullet(d, 'Paste the ID into backend/src/config/smsTemplates.js and set '
              'registered to true. Nothing else changes.',
           bold_prefix='When granted: ')

    # ── 2. Where SMS is used ──
    h(d, '2.  Where the platform sends an SMS')

    para(d, 'Read out of the code rather than from memory; the source of each '
            'row is named so it can be checked.')

    table(d,
          ['#', 'Journey', 'Internal purpose', 'When it is sent', 'Source', 'OTP?'],
          [[a, b, c, dd, e, f] for a, b, c, dd, e, f in USAGE],
          [0.28, 1.5, 1.05, 2.15, 1.35, 0.42],
          font=8.5, mono_cols=(2, 4))

    d.add_paragraph()

    h(d, '2.1  How a code is put into the message', 2, space_before=10)
    para(d, 'The registered wording is stored with its placeholder intact and '
            'filled at send time. All four registrations take a single '
            'variable — the code, or for the security notice the last four '
            'digits of the new number. Jio\'s portal writes the placeholder as '
            '{#number#}; the code writes it as {#var#}. They denote the same '
            'thing, and the text that reaches the carrier — with the value '
            'already substituted — is identical either way.')
    para(d, 'None of the approved wordings mention an expiry period, so none is '
            'sent. The validity of a code is shown on screen instead, where no '
            'carrier has an opinion about it.')
    para(d, 'The message text and the template ID are always sent together, '
            'because the carrier checks the two against each other. This is '
            'the reason the two other registered IDs in section 4 cannot '
            'simply be switched on: sending our wording under their ID would '
            'be dropped exactly as surely as sending unregistered text.',
         muted=True)

    h(d, '2.2  The one message with no template', 2, space_before=12)
    para(d, 'Journey 5 sends free text — it is a notice, not a code:')
    mono(d, 'Your IFQM sign-in number was changed to one ending 1234. '
            'If this was not you, contact your administrator.')
    para(d, 'It is currently sent under the activation template ID, whose '
            'approved wording is completely different, so the carrier will not '
            'deliver it. Template T4 in the next section is the registration '
            'it needs. The number shown is the last four digits of the new '
            'mobile number, which is why T4 carries one variable.')

    d.add_page_break()

    # ── 3. The request ──
    h(d, '3.  Content templates requested', 1, space_before=0)

    para(d, f'All four are to be registered under PE ID {PE_ID} with header '
            f'{HEADER}, in plain English (not Unicode). Each variable is a '
            f'{{#var#}} placeholder; none of them exceeds the 30-character '
            f'limit — the codes are six digits, the validity is a one- or '
            f'two-digit number of minutes, and the T4 variable is four digits.')

    table(d,
          ['Ref', 'Purpose', 'Category', 'Vars', 'Message text as it should be registered'],
          [[a, b, c, dd, e] for a, b, c, dd, e, _ in REQUESTS],
          [0.42, 1.55, 1.0, 0.4, 3.4],
          font=8.5, mono_cols=(4,))

    d.add_paragraph()
    para(d, 'Where each one is used:', bold=True)
    for ref, name, _cat, _v, _txt, used in REQUESTS:
        bullet(d, used, bold_prefix=f'{ref} — {name}: ')

    h(d, '3.1  The exact wording, one per line', 2, space_before=14)
    para(d, 'Reproduced separately so the text can be copied into the DLT '
            'portal without the table formatting coming with it. Punctuation '
            'and spacing matter: the carrier matches the registered text '
            'character for character apart from the variables.', muted=True)
    for ref, name, _cat, nvars, txt, _used in REQUESTS:
        p = d.add_paragraph()
        r = p.add_run(f'{ref} · {name} · {nvars} variable(s)')
        r.bold = True
        r.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        mono(d, txt)

    h(d, '3.2  A note on the category', 2, space_before=12)
    para(d, 'All four are proposed as Service Implicit: they are sent only in '
            'response to something the recipient has just done on the platform '
            '— asking to sign in, asking for a code, changing their own '
            'number — and none of them is promotional. If the operator\'s '
            'portal classifies authentication codes under a separate OTP or '
            'Transactional heading, T1 to T3 should follow that classification '
            'and T4 should remain Service Implicit, since it carries no code. '
            'This is worth confirming with the operator before submitting, as '
            'the available categories differ between DLT portals.')

    # ── 4. Existing registrations ──
    h(d, '4.  Registrations we already hold')

    table(d,
          ['Template ID', 'Named as', 'Approved wording known?', 'Status'],
          [[tid, label, 'Yes', 'Registered and in use'] for _p, label, tid, _t in REGISTERED]
          + [['—', PENDING[1], 'Yes (submitted)', PENDING[3]]],
          [1.55, 1.15, 1.35, 2.7], font=9, mono_cols=(0,))

    d.add_paragraph()
    para(d, 'The two unusable registrations are not a problem with the '
            'registrations themselves. They exist and are valid; what is '
            'missing is the approved text that was registered against them. '
            'Because the carrier checks the ID and the text as a pair, we '
            'cannot send anything under an ID whose registered wording we '
            'cannot reproduce exactly.')

    para(d, 'There are therefore two ways to close this, and either is '
            'acceptable:', bold=True)
    bullet(d, 'Ask the operator for the approved wording already registered '
              'against 1207177450582613311 and 1207177450911544422. If it is '
              'suitable, we adopt it as-is and only T2 and T4 need to be '
              'newly registered.',
           bold_prefix='Option A — recover the wording: ')
    bullet(d, 'Register all four templates in section 3 with the wording given '
              'there, and retire the two older IDs. This is the cleaner '
              'outcome: the wording then matches what each journey actually '
              'does, in language the recipient can act on.',
           bold_prefix='Option B — register afresh: ')

    # ── 5. After approval ──
    h(d, '5.  What changes on our side once the templates are approved')

    para(d, 'T1 to T3 need no change to the application at all. The template '
            'IDs and their wording are configuration, held in the deployment '
            'environment and read at start-up:')

    table(d,
          ['Setting', 'Receives'],
          [['SMS_TEMPLATE_LOGIN / SMS_TEXT_LOGIN', 'T1 — ID and text'],
           ['SMS_TEMPLATE_ACTIVATION / SMS_TEXT_ACTIVATION', 'T2 — ID and text'],
           ['SMS_TEMPLATE_RESET / SMS_TEXT_RESET', 'T3 — ID and text'],
           ['(new) the number-change alert', 'T4 — ID and text']],
          [3.0, 3.6], font=9, mono_cols=(0,))

    d.add_paragraph()
    para(d, 'Each pair must be updated together — the ID and the text that was '
            'approved for it — because they are checked against each other. '
            'Updating one without the other produces messages that the gateway '
            'accepts and the carrier silently drops, which is the failure mode '
            'hardest to notice: nothing appears wrong at either end, and only '
            'the recipient knows the code never came.', muted=True)

    para(d, 'T4 is the one exception, and it is a small one. The number-change '
            'alert is currently sent under the phone-verification purpose, so '
            'once T4 exists it needs a purpose of its own — a single new '
            'setting and the one line that names it. That is a change we make '
            'here, not something the operator needs to know about; it is '
            'recorded so the work is not forgotten when the approvals come '
            'back.')

    para(d, 'Once T1 to T3 are in place, the platform can come off the shared '
            'activation registration one purpose at a time, and each journey '
            'will describe itself correctly. T4 makes the number-change alert '
            'deliverable for the first time.')

    d.save(OUT)
    print('saved', OUT)


if __name__ == '__main__':
    build()
