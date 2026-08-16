# -*- coding: utf-8 -*-
"""
Assemble organized.docx: merge the remaining source documents in, refresh the
figures that have gone out of date, and clean the typography.

    python docs/_assemble.py

What it does, in order:

  1. Refreshes the four figures whose PNGs were regenerated (A-3, A-4, A-7,
     D-4) in place, so the document shows what the software now does rather
     than what it did when the file was first built.
  2. Appends PART E (Technical Manual) and PART F (Project Flows, Phases and
     Timeline) from the two markdown sources. PART C is deliberately skipped —
     it is the test-case document, removed for now and merged back later, and
     taking its letter would mean renumbering everything after it twice.
  3. Removes empty paragraphs and the page breaks that leave blank pages.
  4. Justifies body text.

It is idempotent: run it twice and the second run changes nothing, because each
appended part is keyed by its heading and skipped if already present.
"""
import io
import os
import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(HERE, 'organized.docx')
DIAGRAMS = os.path.join(HERE, 'diagrams')

sys.stdout.reconfigure(encoding='utf-8', errors='replace')


# ── 1. Figures that were regenerated ──────────────────────────────────────
# Keyed by the caption text that sits under each one, because that is the only
# stable handle: the images themselves carry no name once embedded.
REFRESH = {
    'Figure A-3': 'A3_stack.png',
    'Figure A-4': 'A4_modules.png',
    'Figure A-7': 'A7_integration.png',
    'Figure D-4': 'D4_er_master.png',
}


def refresh_figures(d):
    """Swap the embedded image bytes for the freshly rendered ones."""
    from docx.oxml.ns import qn
    paras = d.paragraphs
    done = []
    for i, p in enumerate(paras):
        cap = p.text.strip()
        m = re.match(r'(Figure [A-Z]-\d+)', cap)
        if not m or m.group(1) not in REFRESH:
            continue
        png = os.path.join(DIAGRAMS, REFRESH[m.group(1)])
        if not os.path.exists(png):
            continue
        # The image sits in the paragraph AFTER its caption in this document,
        # not before it — the caption introduces the figure. Looking backwards
        # found the previous figure's image, or nothing at all.
        for fwd in (1, 2, 3):
            if i + fwd >= len(paras):
                break
            host = paras[i + fwd]
            blips = host._element.findall('.//' + qn('a:blip'))
            if not blips:
                continue
            rid = blips[0].get(qn('r:embed'))
            part = d.part.related_parts[rid]
            part._blob = io.open(png, 'rb').read()
            done.append(m.group(1))
            break
    return done


# ── 2. Markdown → Word ────────────────────────────────────────────────────
def _borders(t):
    """
    Hairline borders, set directly.

    The document's only table style carries none, so a table added with it
    arrives as floating text in invisible cells. The existing tables get their
    rules from direct formatting too, so this matches them rather than
    introducing a second convention.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), 'A6A6A6')
        borders.append(e)
    tblPr.append(borders)


def md_blocks(path):
    """
    Split a markdown file into blocks we know how to render.

    Deliberately small: these two documents use headings, paragraphs, bullets,
    fenced code and pipe tables and nothing else, so a full markdown parser
    would be more surface than the job needs.
    """
    text = io.open(path, encoding='utf-8').read().split('\n')
    out, buf, fence = [], [], False
    for line in text:
        if line.startswith('```'):
            if fence:
                out.append(('code', '\n'.join(buf)))
                buf = []
            fence = not fence
            continue
        if fence:
            buf.append(line)
            continue
        h = re.match(r'^(#{1,6})\s+(.*)$', line)
        if h:
            if buf:
                out.append(('para', ' '.join(buf).strip()))
                buf = []
            out.append(('h%d' % len(h.group(1)), h.group(2).strip()))
            continue
        if re.match(r'^\s*[-*]\s+', line):
            if buf and not buf[-1].startswith('\x00'):
                out.append(('para', ' '.join(buf).strip()))
                buf = []
            out.append(('bullet', re.sub(r'^\s*[-*]\s+', '', line).strip()))
            continue
        if line.strip().startswith('|'):
            out.append(('row', line.strip()))
            continue
        if not line.strip():
            if buf:
                out.append(('para', ' '.join(buf).strip()))
                buf = []
            continue
        buf.append(line.strip())
    if buf:
        out.append(('para', ' '.join(buf).strip()))
    return out


INLINE = re.compile(r'(\*\*.+?\*\*|`.+?`|\*.+?\*)')


def add_rich(p, text):
    """Bold, italic and code spans, so the prose does not arrive as one slab."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith('**') and chunk.endswith('**'):
            p.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith('`') and chunk.endswith('`'):
            r = p.add_run(chunk[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
            p.add_run(chunk[1:-1]).italic = True
        else:
            p.add_run(chunk)


def append_markdown(d, path, part_title, base_level=3, part_style='Heading 1'):
    """Append one markdown document as a PART, mapping its headings under it."""
    if any(p.text.strip() == part_title for p in d.paragraphs):
        print('   already present, skipped:', part_title)
        return 0

    if part_style == 'Heading 1':
        d.add_page_break()
    h = d.add_paragraph(part_title, style=part_style)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pending_rows = []

    def flush_rows():
        if not pending_rows:
            return
        rows = [r for r in pending_rows if not re.match(r'^\|[\s:\-|]+\|$', r)]
        cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
        if cells:
            width = max(len(c) for c in cells)
            t = d.add_table(rows=0, cols=width)
            # 'Normal Table' is the only table style this document defines, and
            # it is what all 169 existing tables use. Asking for 'Table Grid'
            # raises rather than falling back, and a new table that looked
            # different from every other one would be worse anyway.
            t.style = 'Normal Table'
            _borders(t)
            for ri, row in enumerate(cells):
                wc = t.add_row().cells
                for ci in range(width):
                    val = row[ci] if ci < len(row) else ''
                    wc[ci].text = re.sub(r'\*\*|`', '', val)
                    for para in wc[ci].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.bold = True
        pending_rows.clear()

    n = 0
    for kind, val in md_blocks(path):
        if kind != 'row':
            flush_rows()
        if kind.startswith('h'):
            lvl = int(kind[1])
            # The file's own H1 is its title, which the PART heading already
            # states — drop it rather than printing the same words twice.
            if lvl == 1 and n == 0:
                n += 1
                continue
            # The document defines Heading 1-4 only; deeper markdown
            # headings flatten onto Heading 4 rather than raising.
            style = 'Heading %d' % min(base_level + lvl - 2, 4)
            d.add_paragraph(val, style=style)
        elif kind == 'para':
            p = d.add_paragraph()
            add_rich(p, val)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == 'bullet':
            # No List Bullet style in this document, and List Paragraph carries
            # no glyph of its own, so the mark is written in — which is what the
            # existing bulleted passages do.
            p = d.add_paragraph(style='List Paragraph')
            p.add_run('•  ')
            add_rich(p, val)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == 'code':
            p = d.add_paragraph()
            r = p.add_run(val)
            r.font.name = 'Consolas'
            r.font.size = Pt(8.5)
        elif kind == 'row':
            pending_rows.append(val)
        n += 1
    flush_rows()
    print('   appended:', part_title)
    return n


# ── 3 & 4. Typography ─────────────────────────────────────────────────────
def tidy(d):
    """
    Drop empty paragraphs and stray page breaks, and justify the body text.

    An empty paragraph carrying an image or a page break is kept — removing
    those is how a diagram disappears and two sections end up on one page.
    """
    from docx.oxml.ns import qn
    removed = 0
    for p in list(d.paragraphs):
        if p.text.strip():
            continue
        el = p._element
        if el.findall('.//' + qn('a:blip')):
            continue                                   # holds a picture
        if el.findall('.//' + qn('w:br')):
            continue                                   # holds a page break
        if el.getparent() is None:
            continue
        el.getparent().remove(el)
        removed += 1

    justified = 0
    for p in d.paragraphs:
        st = (p.style.name or '')
        if st.startswith('Heading') or st in ('Title', 'Caption'):
            continue
        if not p.text.strip():
            continue
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue                                   # captions stay centred
        if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            justified += 1
    return removed, justified


def main():
    d = docx.Document(DOC)
    print('opened:', DOC)

    print('\n1. refreshing figures')
    print('   updated:', ', '.join(refresh_figures(d)) or 'none')

    print('\n2. appending the remaining documents')
    append_markdown(d, os.path.join(HERE, 'TECHNICAL_MANUAL.md'),
                    'PART E — TECHNICAL MANUAL')
    append_markdown(d, os.path.join(HERE, 'PROJECT_FLOWCHART.md'),
                    'PART F — PROJECT FLOWS, PHASES AND TIMELINE')
    # The minutes go INSIDE Part F rather than becoming a part of their own.
    # Part F is where the project-management material lives — phases, timeline,
    # flows — and a decision record belongs beside the plan it changed, not in
    # an appendix after it.
    append_markdown(d, os.path.join(HERE, 'MOM_29Jul2026_Implementation_Status.md'),
                    'Minutes of Meeting — 29 July 2026, and what was implemented',
                    part_style='Heading 2')

    print('\n3. tidying')
    removed, justified = tidy(d)
    print('   empty paragraphs removed:', removed)
    print('   paragraphs justified    :', justified)

    d.save(DOC)
    print('\nsaved:', DOC)
    print('paragraphs now:', len(docx.Document(DOC).paragraphs))


if __name__ == '__main__':
    main()
