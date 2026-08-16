# -*- coding: utf-8 -*-
"""
Make organized.docx presentable: margins, figure sizing, and page breaks.

    python docs/fix_doc_layout.py

The three things that made it look wrong, in the order they matter:

  1. EVERY SECTION HAD ZERO MARGINS. The text column was the full 8.27in of an
     A4 page, so every line ran from one paper edge to the other. Nothing else
     about the typography can look right underneath that — justified text with
     no margin reads as a wall, and it is the first thing an eye notices.

  2. Figures sat at 61-73% of that column, left-aligned, so a diagram drawn to
     be read at full width was shrunk by a third and pushed to one side. The
     text inside them became too small to read, which is what makes a technical
     document useless rather than merely ugly.

  3. Three figures to a page with no separation, so a caption for one sat
     against the artwork of the next.

The cover page is left exactly as it is — it is a designed full-bleed image and
margins would put a white border around artwork that is meant to reach the edge.
"""
import os
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(HERE, 'organized.docx')

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

MARGIN = Inches(0.85)          # a normal business-document margin
TOP_BOTTOM = Inches(0.8)
FULL_BLEED_MIN = Inches(7.9)   # anything this wide is the cover art, not a figure


def set_margins(d):
    """Give every section a margin. The cover keeps its full bleed."""
    changed = 0
    for i, sec in enumerate(d.sections):
        # Section 0 carries the cover image, which is meant to reach the edge.
        if i == 0:
            continue
        if sec.left_margin and sec.left_margin >= MARGIN:
            continue
        sec.left_margin = MARGIN
        sec.right_margin = MARGIN
        sec.top_margin = TOP_BOTTOM
        sec.bottom_margin = TOP_BOTTOM
        changed += 1
    return changed


def column_width(d):
    s = d.sections[-1]
    return s.page_width - s.left_margin - s.right_margin


def size_figures(d, col):
    """
    Scale every NUMBERED FIGURE to the full text column, keeping its aspect
    ratio, and centre it. A diagram shrunk to two thirds is a diagram whose
    labels cannot be read.

    Only figures. Sizing by width alone caught the university logo on the
    front matter and blew it up to fill the page — a crest is not a diagram,
    and the rule has to know the difference. A figure is identified by the
    "Figure X-N" caption sitting immediately before it, which is the same
    handle the caption formatting uses.
    """
    import re
    paras = d.paragraphs
    figure_hosts = set()
    for i, p in enumerate(paras):
        if not re.match(r'Figure [A-Z]-\d+', p.text.strip()):
            continue
        for fwd in (1, 2, 3):
            if i + fwd < len(paras) and paras[i + fwd]._element.findall('.//' + qn('a:blip')):
                figure_hosts.add(id(paras[i + fwd]._element))
                break

    grown = 0
    for p in d.paragraphs:
        if id(p._element) not in figure_hosts:
            continue
        exts = p._element.findall('.//' + qn('wp:extent'))
        if not exts:
            continue
        for ext in exts:
            w = int(ext.get('cx'))
            h = int(ext.get('cy'))
            if w >= FULL_BLEED_MIN:
                continue                      # the cover
            if w <= 0 or h <= 0:
                continue
            scale = col / w
            if abs(scale - 1.0) < 0.02:
                continue
            new_w, new_h = int(w * scale), int(h * scale)
            ext.set('cx', str(new_w))
            ext.set('cy', str(new_h))
            # The shape extent inside the drawing has to move with it, or Word
            # renders the picture at the old size inside a larger frame.
            for xfrm_ext in p._element.findall('.//' + qn('a:ext')):
                xfrm_ext.set('cx', str(new_w))
                xfrm_ext.set('cy', str(new_h))
            grown += 1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return grown


def space_figures(d):
    """
    Put a figure and its caption on the same page, with air around them.

    keep_with_next on the caption is what stops a caption orphaning at the foot
    of one page with its artwork at the top of the next.
    """
    import re
    touched = 0
    paras = d.paragraphs
    for i, p in enumerate(paras):
        if not re.match(r'Figure [A-Z]-\d+', p.text.strip()):
            continue
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
        # air under the artwork that follows
        if i + 1 < len(paras):
            paras[i + 1].paragraph_format.space_after = Pt(14)
            paras[i + 1].paragraph_format.keep_with_next = False
        touched += 1
    return touched


def break_before_parts(d):
    """Each PART starts on a fresh page; so does the acknowledgement."""
    import re
    n = 0
    for p in d.paragraphs:
        t = p.text.strip()
        starts_part = t.startswith('PART ') and p.style.name.startswith('Heading')
        is_ack = t.upper().startswith('ACKNOWLEDGEMENT')
        if not (starts_part or is_ack):
            continue
        if p.paragraph_format.page_break_before:
            continue
        p.paragraph_format.page_break_before = True
        n += 1
    return n


def heading_spacing(d):
    """Headings need room above them or a section reads as one slab."""
    sizes = {'Heading 1': (22, 10), 'Heading 2': (16, 8),
             'Heading 3': (13, 6), 'Heading 4': (11, 5)}
    n = 0
    for p in d.paragraphs:
        st = p.style.name or ''
        if st not in sizes:
            continue
        before, after = sizes[st]
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.keep_with_next = True
        n += 1
    return n


def main():
    d = docx.Document(DOC)
    print('sections given a margin :', set_margins(d))
    col = column_width(d)
    print('text column is now      :', round(col / 914400, 2), 'in')
    print('figures scaled + centred:', size_figures(d, col))
    print('captions tied to artwork:', space_figures(d))
    print('page breaks added       :', break_before_parts(d))
    print('headings spaced         :', heading_spacing(d))
    d.save(DOC)
    print('\nsaved', DOC)


if __name__ == '__main__':
    main()
