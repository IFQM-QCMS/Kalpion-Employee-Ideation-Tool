#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Build "Software Architecture.docx" - the complete architecture and design
document for the IFQM Employee Ideation Tool.

    python docs/gen_architecture_doc.py

House rules, because they were asked for explicitly and are easy to break:

  * The TEXT is black. Word's built-in Heading styles are blue by default, so
    every style used here has its colour forced to black, and blacken_all_styles
    strips colour from the forty-odd styles nothing uses. Table shading stays a
    light grey. Colour appears only inside the diagrams, where it carries
    meaning - and every diagram also reads correctly without it.
  * No emoji, no icons, no decorative characters in the prose.
  * Diagrams are PNGs drawn by arch_drawings.py. Each figure still carries its
    monospace fallback art, and figure() uses it if the drawing is missing, so
    the document builds either way.
  * Plain English. Where a technical term is unavoidable it is explained in the
    sentence that introduces it.

Every quoted figure - tables, endpoints, modules, screens, tests, translation
keys - was counted from the repository for this version, not estimated. The
endpoint counts in section 10 come from docs/api_inventory.py, which reads the
route files. If you change the code and not this document, the document is
wrong; re-run the counts before issuing another version.
"""
import io
import os
import re
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
#  Document metadata
# ---------------------------------------------------------------------------
DOC_TITLE = "Software Architecture and Design Document"
PRODUCT = "IFQM Employee Ideation Tool (EIT)"
VERSION = "1.1"
STATUS = "Reviewed and verified against the code - for final review"
ISSUE_DATE = date.today().strftime("%d %B %Y")
# Writes alongside the other documents in docs/, not into the project root.
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Software Architecture.docx")

BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x59, 0x59, 0x59)

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


# ---------------------------------------------------------------------------
#  Styling helpers
# ---------------------------------------------------------------------------
def setup_styles(doc):
    """Force every style black and set the base fonts."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    sizes = {"Title": 26, "Heading 1": 16, "Heading 2": 13, "Heading 3": 11.5, "Heading 4": 10.5}
    for name, size in sizes.items():
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = BLACK          # Word defaults these to blue
        st.font.bold = True
        st.font.italic = False
        pf = st.paragraph_format
        pf.space_before = Pt(14 if name == "Heading 1" else 10)
        pf.space_after = Pt(5)
        pf.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        try:
            s = doc.styles[name]
            s.font.name = BODY_FONT
            s.font.size = Pt(10.5)
            s.font.color.rgb = BLACK
        except KeyError:
            pass


def blacken_all_styles(doc):
    """
    Strip colour from EVERY style definition in the document.

    Overriding the handful of styles this document uses is not enough. The
    template python-docx ships with defines about forty more - Heading 5 to 9,
    Subtitle, Caption, Intense Quote, the table styles - and every one of them
    carries a blue or red from the default Office theme. Nothing here uses them,
    so the document renders black either way, but the colours are still sitting
    in the file. Anyone who later applies one of those styles, or inspects the
    XML, would find colour in a document that was asked to have none.

    This removes the colour element from all of them, so "everything is black"
    is true of the file and not only of what happens to be on the page.
    """
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    removed = 0
    for style in doc.styles.element:
        for rpr in style.iter(ns + "rPr"):
            for col in list(rpr.findall(ns + "color")):
                val = (col.get(ns + "val") or "").upper()
                if val not in ("000000", "595959", "AUTO"):
                    rpr.remove(col)
                    removed += 1
    return removed


def shade(cell, hex_fill="EDEDED"):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def h(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLACK
    return p


def para(doc, text, italic=False, size=10.5, space_after=6, align=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = GREY if italic else BLACK
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullets(doc, items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(it)
        run.font.name = BODY_FONT
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
        p.paragraph_format.space_after = Pt(2)


# Figure number → the file drawn for it by arch_drawings.py. A figure with no
# entry here falls back to its monospace art, so adding a drawing is additive
# and removing one cannot break the build.
FIGURE_IMAGES = {
    "Figure A-1": "A1_context", "Figure A-2": "A2_containers",
    "Figure A-3": "A3_stack", "Figure A-4": "A4_modules",
    "Figure A-5": "A5_deployment", "Figure A-6": "A6_security",
    "Figure A-7": "A7_integration",
    "Figure D-1": "D1_workflow", "Figure D-2": "D2_dfd0", "Figure D-3": "D3_dfd1",
    "Figure D-4": "D4_er_master", "Figure D-5": "D5_er_tenant",
    "Figure D-6": "D6_usecase", "Figure D-7": "D7_seq_login",
    "Figure D-8": "D8_seq_submit", "Figure D-9": "D9_seq_approval",
    "Figure D-10": "D10_seq_register", "Figure D-11": "D11_seq_qcms",
    "Figure D-12": "D12_class", "Figure D-13": "D13_screenflow",
    "Figure D-14": "D14_wireframes", "Figure D-15": "D15_errors",
}

DIAGRAM_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "diagrams")


def figure(doc, number, title, art, note=None, legend=None):
    """
    A titled figure with an optional legend and note.

    Every figure is now a drawing. The diagrams used to be box-drawing
    characters in a monospace block, which kept the whole document in one text
    file but read as a wall of dashes — and an entity relationship diagram drawn
    in `+---+` is genuinely hard to follow, because the one thing it cannot show
    is which end of a relationship is the "many" end.

    The monospace art is still passed in and is still used if a drawing is
    missing, so the document builds either way.
    """
    cap = doc.add_paragraph()
    r = cap.add_run("%s  %s" % (number, title))
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True

    image = FIGURE_IMAGES.get(number)
    path = os.path.join(DIAGRAM_DIR, image + ".png") if image else None

    if path and os.path.exists(path):
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_after = Pt(2)
        # Sized to the text column so a wide diagram is never cropped by the
        # page margin. Word scales the height to match.
        pic.add_run().add_picture(path, width=Inches(6.3))
    else:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.text = ""
        for i, line in enumerate(art.rstrip("\n").split("\n")):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            run = para.add_run(line)
            run.font.name = MONO_FONT
            run.font.size = Pt(7.6)
            run.font.color.rgb = BLACK
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run("Legend:  " + legend)
        lr.font.name = BODY_FONT
        lr.font.size = Pt(8.5)
        lr.font.color.rgb = GREY
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(2)
    if note:
        np = doc.add_paragraph()
        nr = np.add_run(note)
        nr.font.name = BODY_FONT
        nr.font.size = Pt(8.5)
        nr.italic = True
        nr.font.color.rgb = GREY
        np.paragraph_format.space_after = Pt(10)


def table(doc, headers, rows, widths=None, font_size=9, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.name = BODY_FONT
        cr.font.size = Pt(10)
        cr.font.color.rgb = BLACK
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after = Pt(4)
        cp.paragraph_format.keep_with_next = True

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(htxt)
        r.bold = True
        r.font.name = BODY_FONT
        r.font.size = Pt(font_size)
        r.font.color.rgb = BLACK
        p.paragraph_format.space_after = Pt(1)
        shade(c)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = BODY_FONT
            r.font.size = Pt(font_size)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(1)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _small(p, text):
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    return r


def _field(p, instr):
    """
    A Word field, so the number is computed when the document is opened.

    The previous version of this built both PAGE and NUMPAGES from one loop and
    decided which was which by inspecting the paragraph text, which had not been
    written yet - so every field came out as PAGE and the footer read "Page 4 of
    4" on every page. Passing the instruction in explicitly removes the guess.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    for tag, val in (("w:rFonts", BODY_FONT), ("w:sz", "16"), ("w:color", "595959")):
        el = OxmlElement(tag)
        el.set(qn("w:ascii") if tag == "w:rFonts" else qn("w:val"), val)
        if tag == "w:rFonts":
            el.set(qn("w:hAnsi"), val)
        rpr.append(el)
    run.append(rpr)
    fld.append(run)
    p._p.append(fld)


def footer_text(doc, text):
    """Document identity on the left, page N of M on the right."""
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # One right-aligned tab stop at the text margin puts the page number on
        # the right edge without a second paragraph or a table.
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
        _small(p, text)
        _small(p, "\t")
        _small(p, "Page ")
        _field(p, "PAGE")
        _small(p, " of ")
        _field(p, "NUMPAGES")


# ---------------------------------------------------------------------------
#  Front matter
# ---------------------------------------------------------------------------
def front_matter(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(DOC_TITLE)
    r.font.name = BODY_FONT
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = BLACK
    t.paragraph_format.space_before = Pt(120)
    t.paragraph_format.space_after = Pt(6)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(PRODUCT)
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    r.font.color.rgb = BLACK
    s.paragraph_format.space_after = Pt(40)

    table(doc, ["Field", "Detail"], [
        ["Document title", DOC_TITLE],
        ["Product", PRODUCT],
        ["Version", VERSION],
        ["Status", STATUS],
        ["Date of issue", ISSUE_DATE],
        ["Prepared by", "Development team"],
        ["Reviewed by", "(pending)"],
        ["Approved by", "(pending)"],
        ["Classification", "Internal - for IFQM and project stakeholders"],
    ], widths=[1.9, 4.3], font_size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    table(doc, ["Version", "Date", "Author", "Summary of change"], [
        ["1.0", "02 August 2026", "Development team",
         "First complete version. Covered all 25 requested deliverables and described the "
         "system as built."],
        ["1.1", ISSUE_DATE, "Development team",
         "Review pass. Section 10 rewritten with per-endpoint request, response, "
         "authentication and error detail. Section 13 restated as module structure rather than "
         "a class diagram, since the codebase has no classes. New section 26 covering work not "
         "yet done. Every diagram, count and traceability row checked against the code: "
         "endpoint, table, module, screen and test figures corrected, and the requirements "
         "matrix extended to cover subscriptions and billing, which the previous version listed "
         "as out of scope after it had been built."],
    ], widths=[0.7, 1.05, 1.15, 3.3], font_size=9,
        caption="Revision history")

    doc.add_page_break()

    h(doc, "How to read this document", 1)
    para(doc, "This document describes how the IFQM Employee Ideation Tool is built. It is "
              "written for a mixed audience, so it avoids jargon wherever it can and explains "
              "it where it cannot. If you only need the shape of the system, sections 1 to 6 "
              "are enough. If you are going to work on it, read the whole thing.")

    h(doc, "The way it is organised", 2)
    table(doc, ["Part", "What it covers", "Who it is mainly for"], [
        ["A", "Overview and requirements (sections 1-2)", "Everybody"],
        ["B", "Architecture diagrams (sections 3-6) - what the system is made of and where each part runs",
         "Anyone reviewing the design"],
        ["C", "Detailed design diagrams (sections 7-14) - how it behaves, step by step",
         "Developers and testers"],
        ["D", "Running it (sections 15-23) - deployment, security, failure, monitoring, testing, release",
         "Operations and QA"],
        ["E", "Decisions, traceability and outstanding work (sections 24-26)",
         "Reviewers and auditors"],
    ], widths=[0.5, 4.2, 1.6], font_size=9)

    h(doc, "Architecture diagrams and design diagrams", 2)
    para(doc, "The two are kept apart deliberately, because they answer different questions. An "
              "architecture diagram shows the pieces and how they are arranged; it would still "
              "be true if nobody ever pressed a button. A design diagram shows what actually "
              "happens when somebody does. Mixing the two is the usual reason architecture "
              "documents become unreadable.")
    para(doc, "The separation is carried in the figure number itself, so it holds wherever a "
              "figure appears:")
    table(doc, ["Prefix", "Kind", "Answers", "Figures"], [
        ["A-n", "Architecture - structure", "What is the system made of, and where does each "
                                            "piece run?", "A-1 to A-7"],
        ["D-n", "Detailed design - behaviour", "What happens, step by step, when somebody does "
                                               "something?", "D-1 to D-15"],
    ], widths=[0.6, 1.6, 2.6, 1.4], font_size=9)
    para(doc, "Part B holds the architecture set and Part C the design set. Three architecture "
              "figures (A-5 deployment, A-6 security, A-7 integrations) and one design figure "
              "(D-15 failure handling) appear later, in Part D, because each belongs beside the "
              "operational discussion that uses it - a deployment diagram is of little use "
              "eleven pages away from the environment table. The prefix still says which kind "
              "each one is, and the list of figures below names the kind and the part for every "
              "figure in the document.")

    h(doc, "Conventions used in the diagrams", 2)
    bullets(doc, [
        "Every diagram is a drawing, generated from a script rather than placed by hand, so a "
        "diagram and the thing it describes are changed in one action.",
        "Colour carries meaning and is consistent across every figure: indigo for the system "
        "and the path through it, teal for stored data, amber for a decision, green for a "
        "wanted outcome, red for a refusal, slate for anything outside our control.",
        "Colour is never the only signal. Every diagram reads correctly in black and white, and "
        "each one carries a legend directly beneath it.",
        "A solid box is something that exists. A solid arrow is something that happens now.",
        "Anything marked optional works only if that organisation has configured it. The "
        "product runs without any of them.",
        "Where a number is quoted - tables, endpoints, modules, screens, tests - it was counted "
        "from the code for this version, not estimated.",
    ])

    h(doc, "Assumptions this document rests on", 2)
    table(doc, ["#", "Assumption", "If it turns out to be wrong"], [
        ["A1", "Customers are small and medium businesses, mostly manufacturing, of roughly "
               "20 to 500 staff.",
         "Scale figures in section 20 would need revisiting, though nothing structural changes."],
        ["A2", "Each customer wants their data kept physically separate from other customers.",
         "The separate-database design could be simplified, but this is a common condition of sale."],
        ["A3", "Ideas are typed by people. There is no bulk import of ideas and no machine feed.",
         "The submission path would need a different entry point."],
        ["A4", "Traffic is bursty and modest - busy at the start of a campaign, quiet between.",
         "Sustained heavy traffic would need the caching described in section 20."],
        ["A5", "IFQM operates the platform and holds the customer relationship.",
         "The platform console assumes exactly this separation."],
        ["A6", "Email and SMS are optional per customer. Many will run without either.",
         "Nothing breaks; the features they power simply stay off."],
        ["A7", "Idea text is commercially sensitive to the customer.",
         "The visibility rules in section 16 could be relaxed, and are already a setting."],
    ], widths=[0.4, 3.1, 2.8], font_size=9)

    h(doc, "What this document does not cover", 2)
    bullets(doc, [
        "Taking payment. Plans, GST pricing, trials and holds are all built and are described "
        "in sections 2.1, 10.6.10 and 25. Connecting a payment gateway and issuing invoices is "
        "not, and is not designed here. See section 26.3.",
        "Commercial terms - what IFQM actually charges a given customer. The system records "
        "whichever figure the platform team sets.",
        "Single sign-on with the QCMS, DWM and Skills tools. Agreed in principle, blocked on "
        "an Azure tenant that does not exist yet. See ADR-011.",
        "The manual test case catalogue and the user guide, which are separate documents.",
    ])
    para(doc, "Section 26 lists everything that is not yet done, including the items above. It "
              "is deliberately the last section rather than a footnote, because it is the part "
              "of a handover most likely to be needed.")


# ---------------------------------------------------------------------------
#  Table of contents
# ---------------------------------------------------------------------------
def contents(doc):
    doc.add_page_break()
    h(doc, "Contents", 1)
    para(doc, "Word can generate a live table of contents with page numbers: place the cursor "
              "here and use References, then Table of Contents. The list below is the section "
              "order.", italic=True)

    rows = [
        ["Part A", "Overview and requirements", ""],
        ["", "1", "System Overview and Objectives"],
        ["", "2", "Functional and Non-Functional Requirements"],
        ["Part B", "Architecture diagrams", ""],
        ["", "3", "Overall System Architecture"],
        ["", "4", "Architecture Description"],
        ["", "5", "Technology Stack"],
        ["", "6", "Module and Component Structure"],
        ["Part C", "Detailed design diagrams", ""],
        ["", "7", "System Workflow"],
        ["", "8", "Data Flow"],
        ["", "9", "Database Design"],
        ["", "10", "API Architecture and Specification"],
        ["", "11", "Use Cases"],
        ["", "12", "Sequence Diagrams"],
        ["", "13", "Module and Service Structure"],
        ["", "14", "User Interface and Screen Flow"],
        ["Part D", "Running the system", ""],
        ["", "15", "Deployment Architecture"],
        ["", "16", "Security Architecture"],
        ["", "17", "Integration Architecture"],
        ["", "18", "Error Handling and Exception Flow"],
        ["", "19", "Logging and Monitoring"],
        ["", "20", "Performance and Scalability"],
        ["", "21", "Backup and Disaster Recovery"],
        ["", "22", "Testing Strategy"],
        ["", "23", "Deployment and Release Plan"],
        ["Part E", "Decisions, traceability and outstanding work", ""],
        ["", "24", "Architecture Decision Records"],
        ["", "25", "Requirements Traceability Matrix"],
        ["", "26", "Work Not Yet Done"],
    ]
    table(doc, ["Part", "No.", "Section"], rows, widths=[0.8, 0.5, 4.9], font_size=9,
          caption="Contents")

    h(doc, "List of figures", 2)
    para(doc, "Twenty-two figures. Architecture figures are numbered A-n and describe structure; "
              "design figures are numbered D-n and describe behaviour. The kind column is the "
              "same distinction, spelled out.")
    table(doc, ["Figure", "Title", "Kind", "Part", "Section"], [
        ["A-1", "System context", "Architecture", "B", "3.1"],
        ["A-2", "Containers - what runs where", "Architecture", "B", "3.2"],
        ["A-3", "Technology stack", "Architecture", "B", "5"],
        ["A-4", "Modules and their dependencies", "Architecture", "B", "6"],
        ["A-5", "Deployment", "Architecture", "D", "15"],
        ["A-6", "Security layers", "Architecture", "D", "16"],
        ["A-7", "Integrations", "Architecture", "D", "17"],
        ["D-1", "Idea lifecycle", "Design - workflow", "C", "7"],
        ["D-2", "Data flow, level 0 (context)", "Design - data flow", "C", "8"],
        ["D-3", "Data flow, level 1 (inside the system)", "Design - data flow", "C", "8"],
        ["D-4", "Registry tables and relationships", "Design - entity relationship", "C", "9.1"],
        ["D-5", "Per-organisation tables and relationships", "Design - entity relationship",
         "C", "9.2"],
        ["D-6", "Use case overview", "Design - use case", "C", "11"],
        ["D-7", "Sign in with a password", "Design - sequence", "C", "12.1"],
        ["D-8", "Submit an idea", "Design - sequence", "C", "12.2"],
        ["D-9", "A decision, and what happens when nobody decides", "Design - sequence",
         "C", "12.3"],
        ["D-10", "MSME registration and approval", "Design - sequence", "C", "12.4"],
        ["D-11", "Push an approved idea to QCMS", "Design - sequence", "C", "12.5"],
        ["D-12", "Module structure, taking the idea path as the example", "Design - module",
         "C", "13"],
        ["D-13", "Screen flow", "Design - UI flow", "C", "14"],
        ["D-14", "Wireframes of the main screens", "Design - wireframe", "C", "14"],
        ["D-15", "How a failure is handled", "Design - exception flow", "D", "18"],
    ], widths=[0.55, 2.75, 1.45, 0.4, 0.65], font_size=8.5,
        caption="List of figures")

    para(doc, "Every diagram type asked for is present: architecture (A-1, A-2, A-3), module "
              "(A-4, D-12), data flow (D-2, D-3), entity relationship (D-4, D-5), use case "
              "(D-6), sequence (D-7 to D-11), deployment (A-5), security (A-6), integration "
              "(A-7), and interface flow and wireframes (D-13, D-14). There is no class diagram, "
              "and section 13 explains why rather than leaving a reader to wonder.", italic=True)

    h(doc, "List of tables", 2)
    para(doc, "Most tables in this document run inside the prose they belong to. These are the "
              "ones worth finding on their own.")
    table(doc, ["Table", "Title", "Section"], [
        ["2.1", "Functional requirements", "2.1"],
        ["2.2", "Non-functional requirements", "2.2"],
        ["6.1", "All 28 screens, grouped by who reaches them", "6.2"],
        ["10.1", "Every route group, with its endpoint count", "10.5"],
        ["13.1", "All 35 service modules", "13.2"],
        ["25.1", "Requirements traceability - functional", "25"],
        ["25.2", "Requirements traceability - non-functional", "25"],
    ], widths=[0.7, 4.2, 0.8], font_size=9)


# ---------------------------------------------------------------------------
#  Assemble
# ---------------------------------------------------------------------------
def main():
    import sys
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    import arch_sections_a, arch_sections_b, arch_sections_c
    import arch_sections_d, arch_sections_e

    doc = Document()
    setup_styles(doc)
    stripped = blacken_all_styles(doc)

    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.7)

    front_matter(doc)
    contents(doc)

    api = (h, para, bullets, table, figure, None)
    for mod in (arch_sections_a, arch_sections_b, arch_sections_c,
                arch_sections_d, arch_sections_e):
        mod.build(doc, *api)

    footer_text(doc, "%s  |  %s  |  Version %s  |  %s" % (PRODUCT, DOC_TITLE, VERSION, ISSUE_DATE))

    core = doc.core_properties
    core.title = DOC_TITLE
    core.subject = PRODUCT
    core.author = "Development team"
    core.comments = STATUS

    out = os.path.abspath(OUT)
    doc.save(out)

    drawn = sum(1 for n in FIGURE_IMAGES.values()
                if os.path.exists(os.path.join(DIAGRAM_DIR, n + ".png")))
    print("Written: %s" % out)
    print("  paragraphs : %d" % len(doc.paragraphs))
    print("  tables     : %d" % len(doc.tables))
    print("  figures    : %d of %d drawn (the rest would fall back to monospace art)"
          % (drawn, len(FIGURE_IMAGES)))
    print("  colour removed from %d unused style definitions" % stripped)


if __name__ == "__main__":
    main()
